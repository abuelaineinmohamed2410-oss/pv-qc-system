import streamlit as st
import pdfplumber
from docx import Document

# =========================
# CONFIG
# =========================
st.set_page_config(page_title="PV QC System", layout="wide")

st.title("Pharmacovigilance Control System")
st.markdown("QC vs Agent – robust pharma-grade mismatch detection")

# =========================
# CLEAN TEXT
# =========================
def clean(text):
    if not text:
        return ""
    text = text.replace("\xa0", " ")
    text = text.replace("", "\n")
    text = text.replace("•", "\n")
    return text

# =========================
# PDF READER
# =========================
def read_pdf(file):
    text = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text)

# =========================
# DOCX READER (FIXED)
# =========================
def read_docx(file):
    doc = Document(file)

    text = []

    # paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text.strip())

    # tables
    for table in doc.tables:
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                if cell.text.strip():
                    row_cells.append(cell.text.strip())
            if row_cells:
                text.append(" ".join(row_cells))

    return "\n".join(text)

# =========================
# FILE HANDLER
# =========================
def extract_text(file):
    if file.name.lower().endswith(".pdf"):
        return read_pdf(file)
    return read_docx(file)

# =========================
# SMART PV FIELD PARSER (FINAL FIX)
# =========================
def extract_fields(text):

    text = clean(text)
    text = text.replace("", "\n").replace("•", "\n")

    lines = [l.strip() for l in text.split("\n") if l.strip()]

    fields = {
        "drug": "",
        "dose": "",
        "indication": "",
        "mrd": "",
        "patient_id": "",
        "dob": "",
        "age": "",
        "gender": "",
        "country": ""
    }

    def get_value(i):
        """
        Collects value from same line or next lines (fix for broken Word/PDF)
        """
        line = lines[i]

        # same line value
        if ":" in line:
            value = line.split(":", 1)[1].strip()
            if value:
                return value

        # multi-line value
        value_parts = []
        for j in range(i + 1, min(i + 4, len(lines))):
            if ":" in lines[j]:
                break
            value_parts.append(lines[j])

        return " ".join(value_parts).strip()

    for i, line in enumerate(lines):

        l = line.lower()

        # DRUG (CRITICAL FIX)
        if "product name" in l:
            fields["drug"] = get_value(i)

        elif "dose" in l:
            fields["dose"] = get_value(i)

        elif "indication" in l:
            fields["indication"] = get_value(i)

        elif "date reported" in l or "mrd" in l:
            fields["mrd"] = get_value(i)

        elif "patient id" in l:
            fields["patient_id"] = get_value(i)

        elif "date of birth" in l:
            fields["dob"] = get_value(i)

        elif "age" in l:
            fields["age"] = get_value(i)

        elif "gender" in l:
            fields["gender"] = get_value(i)

        elif "country" in l:
            fields["country"] = get_value(i)

    return {k: v.lower().strip() for k, v in fields.items() if v}

# =========================
# SEVERITY ENGINE
# =========================
def severity(field):

    if field in ["mrd", "dob", "patient_id"]:
        return "HIGH"

    if field in ["drug", "dose", "indication"]:
        return "MODERATE"

    return "LOW"

# =========================
# COMPARE ENGINE
# =========================
def compare(qc, agent):

    results = []

    keys = set(qc.keys()).union(set(agent.keys()))

    for k in keys:

        qc_val = qc.get(k, "MISSING")
        ag_val = agent.get(k, "MISSING")

        if qc_val != ag_val:

            results.append({
                "field": k.upper(),
                "qc": qc_val,
                "agent": ag_val,
                "severity": severity(k)
            })

    return results

# =========================
# UI
# =========================
qc_file = st.file_uploader("Upload QC File (PDF / DOCX)", type=["pdf", "docx"])
agent_file = st.file_uploader("Upload Agent File (PDF / DOCX)", type=["pdf", "docx"])

if qc_file and agent_file:

    qc_text = extract_text(qc_file)
    agent_text = extract_text(agent_file)

    st.subheader("Raw QC Text")
    st.text_area("", qc_text, height=200)

    st.subheader("Raw Agent Text")
    st.text_area("", agent_text, height=200)

    qc_data = extract_fields(qc_text)
    agent_data = extract_fields(agent_text)

    st.subheader("QC Extracted Data")
    st.json(qc_data)

    st.subheader("Agent Extracted Data")
    st.json(agent_data)

    if st.button("Run QC Validation"):

        results = compare(qc_data, agent_data)

        st.subheader("Mismatch Report")

        if not results:
            st.success("No discrepancies detected")
        else:
            for r in results:

                st.markdown(f"""
### {r['field']}

QC Value: **{r['qc']}**  
Agent Value: **{r['agent']}**  
Severity: **{r['severity']}**

---
""")
